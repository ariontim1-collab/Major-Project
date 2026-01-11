from tkinter import filedialog, simpledialog
import tkinter as tk
from argparse import ArgumentParser
import os
import sys
import openpyxl
from PyPDF2 import PdfReader
from docx import Document
import pandas as pd
from urllib.parse import quote
import comtypes.client
from openpyxl.worksheet.table import Table, TableStyleInfo, TableColumn

def extract_specific_fields(pdf_file, desired_fields):
    """Extract specific form fields from a PDF"""
    pdf = PdfReader(pdf_file)
    field_values = []

    for page_num, page in enumerate(pdf.pages, start=1):
        data = {}
        if "/Annots" in page:
            annotations = page["/Annots"]
            for annotation in annotations:
                obj = pdf.get_object(annotation)
                if obj.get("/FT") == "/Tx":  # Check if it's a text field
                    field_name = obj.get("/T") or obj.get("/TU")  # Field name variation
                    field_value = obj.get("/V")
                    if field_name in desired_fields:
                        data[field_name] = field_value

            if data:
                field_values.append(data)

    print(f"Debug: field_values for page {page_num}: {field_values}")  # Debug information
    return field_values

def write_to_excel(all_fields, field_values, sheet, next_row):
    """Write extracted form fields to an Excel sheet"""

    # Extract the PDF titles as fields
    pdf_titles = all_fields

    # Write titles (field names) if the table doesn't exist
    title_row = next_row
    for idx, field_name in enumerate(pdf_titles, start=1):
        cell = sheet.cell(row=title_row, column=idx, value=field_name)
        cell.font = openpyxl.styles.Font(bold=True)
        cell.fill = openpyxl.styles.PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')

    next_row += 1  # Increment the row counter

    # Concatenate the values of each field in the same row
    row_data = [''] * len(pdf_titles)
    for page_data in field_values:
        for field_name, field_value in page_data.items():
            if field_name in pdf_titles:
                idx = pdf_titles.index(field_name)
                row_data[idx] += str(field_value)  # Concatenate the values without newline character

    # Append the concatenated row data to the sheet
    sheet.append(row_data)
    next_row += 1  # Increment the row counter

    # Create a new table with the user-defined name
    table_range = f"A2:{openpyxl.utils.get_column_letter(len(pdf_titles))}{next_row - 1}"
    table = openpyxl.worksheet.table.Table(displayName=sheet.title, ref=table_range)
    style = openpyxl.worksheet.table.TableStyleInfo(
        name="TableStyleMedium9", showFirstColumn=False,
        showLastColumn=False, showRowStripes=True, showColumnStripes=True
    )
    table.tableStyleInfo = style
    sheet.add_table(table)

    return next_row  # Return the row number for the next data set

def process_pdf_forms(pdf_files, output_file, table_name, pdf_fields):
    """Process multiple PDF forms with user-defined table name and PDF fields"""
    if os.path.exists(output_file):
        workbook = openpyxl.load_workbook(output_file)
    else:
        workbook = openpyxl.Workbook()

    sheet = workbook.create_sheet(title=table_name)  # Create a new sheet with the specified name

    existing_row = 1 if not sheet.max_row else sheet.max_row + 1  # Determine the next row

    for pdf_file in pdf_files:
        field_values = extract_specific_fields(pdf_file, pdf_fields)
        existing_row = write_to_excel(pdf_fields, field_values, sheet, existing_row)

    # Adjust column widths for better readability
    for col in range(1, len(pdf_fields) + 1):
        col_letter = openpyxl.utils.get_column_letter(col)
        sheet.column_dimensions[col_letter].width = 30  # Set column width to 30 (you can adjust this value as needed)

    workbook.save(output_file)
    print(f"Form contents from multiple PDFs saved to '{output_file}' with table '{table_name}'")

def parse_cli():
    """Load command line arguments"""
    parser = ArgumentParser(description='Dump the form contents of PDFs into a single Excel file.')
    parser.add_argument('files', metavar='pdf_form', nargs='+',
                        help='PDF Forms to dump the contents of')
    parser.add_argument('-o', '--out', help='Write output to Excel file',
                        default='output1.xlsx', metavar='FILE')
    return parser.parse_args()

def open_excel_file(file_path):
    """Open the Excel file"""
    if sys.platform.startswith('darwin'):  # For MacOS
        os.system('open ' + file_path)
    elif os.name == 'nt':  # For Windows
        os.startfile(file_path)
    elif os.name == 'posix':  # For Linux
        os.system('xdg-open ' + file_path)
    else:
        print("Unsupported operating system. Couldn't open the file.")

def replace_placeholder_spanning_runs(paragraph, data):
    # Combine the text of all runs to check for placeholders
    full_text = ''.join(run.text for run in paragraph.runs)
    
    # Track if any replacement is made
    replacements_made = False
    
    # Replace all placeholders found in the combined text
    for key, value in data.items():
        placeholder = '{{' + key + '}}'
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, str(value))
            replacements_made = True
    
    # If replacements were made, we need to update the runs
    if replacements_made:
        # Clear all existing runs
        for run in paragraph.runs:
            run.clear()
        # Add a new run with the updated full text
        paragraph.add_run(full_text)

def replace_placeholder_in_tables(tables, data):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholder_spanning_runs(paragraph, data)

# Function to replace placeholders throughout the entire document
def replace_all_placeholders_entire_document(doc, data):
    for para in doc.paragraphs:
        replace_placeholder_spanning_runs(para, data)
    
    # Process tables
    replace_placeholder_in_tables(doc.tables, data)

    # Search in headers and footers across all sections
    for section in doc.sections:
        for header in section.header.paragraphs:
            replace_placeholder_spanning_runs(header, data)
        for footer in section.footer.paragraphs:
            replace_placeholder_spanning_runs(footer, data)
        if section.first_page_footer:
            for first_page_footer in section.first_page_footer.paragraphs:
                replace_placeholder_spanning_runs(first_page_footer, data)
        if section.even_page_footer:
            for even_page_footer in section.even_page_footer.paragraphs:
                replace_placeholder_spanning_runs(even_page_footer, data)


# New function to ask for user input for row selection
def select_row():
    row_num = int(input("Enter the row number to use for data replacement: "))
    return row_num

# Function to save Word document as PDF using comtypes
def save_as_pdf(input_path, output_path):
    word = comtypes.client.CreateObject('Word.Application',dynamic=True)
    # input_path = input_path.replace(' ','\ ')
    # Ensure full path to Word file is provided
    input_path = os.path.abspath(input_path)
    # Encode special characters in the input path
    #encoded_input_path = quote(input_path)
    print(input_path)
    doc = word.Documents.Open(input_path)
    print(doc)
    doc.SaveAs(output_path)#, FileFormat=17)  # 17 represents wdFormatPDF
    doc.Close()
    word.Quit()


def main():
    root = tk.Tk()
    root.withdraw()  # Hide the main window

    # Step 2: Select PDF file
    pdf_file_path = filedialog.askopenfilename(title="Select IDF PDF File", filetypes=[("PDF Files", "*.pdf")])
    if not pdf_file_path:
        print("No PDF file selected.")
        return

    # Step 3: Enter table name
    table_name = simpledialog.askstring("Table Name", "Enter the name for the table:")
    if not table_name:
        print("Table name is required.")
        return

    # Step 4: Enter PDF fields
    pdf_fields_input = simpledialog.askstring("PDF Fields", "Enter PDF fields to extract (comma-separated):")
    if not pdf_fields_input:
        print("PDF fields are required.")
        return
    pdf_fields = [field.strip() for field in pdf_fields_input.split(',')]

    # Process PDF and write to Excel
    output_excel_path = 'output1.xlsx'
    process_pdf_forms([pdf_file_path], output_excel_path, table_name, pdf_fields)

    # Step 6: Select Word document
    word_file_path = filedialog.askopenfilename(title="Select AA Word Document", filetypes=[("Word Files", "*.docx")])
    if not word_file_path:
        print("No Word document selected.")
        return

    # Step 5: Use the last row of data for replacement
    data_df = pd.read_excel(output_excel_path)
    if data_df.empty:
        print("The Excel file is empty.")
        return
    test_data = data_df.iloc[-1].to_dict()

    # Replace placeholders in Word document
    doc = Document(word_file_path)
    replace_all_placeholders_entire_document(doc, test_data)

    # Step 7: Save temporary Word document
    temp_document_path = filedialog.asksaveasfilename(defaultextension=".docx", title="Save Temporary AA Word File", filetypes=[("Word Files", "*.docx")])
    if not temp_document_path:
        print("No file path provided for saving the temporary Word document.")
        return
    doc.save(temp_document_path)

    # Step 8: Save as PDF
    pdf_document_path = filedialog.asksaveasfilename(defaultextension=".pdf", title="Save PDF File", filetypes=[("PDF Files", "*.pdf")])
    if not pdf_document_path:
        print("No file path provided for saving the PDF document.")
        return
    save_as_pdf(temp_document_path, pdf_document_path)

    # Step 9: Delete temporary Word document
    os.remove(temp_document_path)

if __name__ == '__main__':
    main()